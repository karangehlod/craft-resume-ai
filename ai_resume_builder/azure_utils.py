from azure.ai.textanalytics import TextAnalyticsClient
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import ServiceRequestError
from azure.storage.blob import BlobServiceClient
from docx import Document
from dotenv import load_dotenv
import os
import logging

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def authenticate_client():
    key = os.getenv("AZURE_TEXT_ANALYTICS_KEY")
    print(os.getenv("AZURE_TEXT_ANALYTICS_KEY"))
    print(os.getenv("AZURE_ENDPOINT"))
    endpoint = os.getenv("AZURE_ENDPOINT")
    ta_credential = AzureKeyCredential(key)
    text_analytics_client = TextAnalyticsClient(
        endpoint=endpoint, credential=ta_credential)
    return text_analytics_client

def extract_key_phrases(text):
    client = authenticate_client()
    try:
        response = client.extract_key_phrases(documents=[text])
        return response[0].key_phrases
    except ServiceRequestError as e:
        logging.error(f"Service request error: {e}")
        return []

def extract_key_phrases_in_batches(text, batch_size=5120):
    client = authenticate_client()
    text_chunks = [text[i:i + batch_size] for i in range(0, len(text), batch_size)]
    key_phrases = []
    for chunk in text_chunks:
        response = client.extract_key_phrases(documents=[{"id": "1", "language": "en", "text": chunk}])[0]
        if hasattr(response, 'key_phrases'):
            key_phrases.extend(response.key_phrases)
    return key_phrases

def text_similarity(text1, text2):
    client = authenticate_client()
    documents = [{"id": "1", "text": text1}, {"id": "2", "text": text2}]
    response = client.analyze_sentiment(documents=documents)
    return response[0].confidence_scores.positive

def identify_missing_skills(job_skills, resume_skills):
    logging.info(f"Identifying missing skills from job skills: {job_skills} and resume skills: {resume_skills}")
    return [skill for skill in job_skills if skill not in resume_skills]

def generate_resume_docx(user_details, job_skills):
    logging.info(f"Generating resume docx for user details: {user_details} and job skills: {job_skills}")
    doc = Document()
    doc.add_heading('Resume', 0)
    
    doc.add_heading('User Details', level=1)
    doc.add_paragraph(user_details)
    
    doc.add_heading('Job Skills', level=1)
    doc.add_paragraph(job_skills)
    
    doc.save('resume.docx')
    return {"message": "Resume generated successfully"}

def upload_to_blob(file_path, container_name):
    logging.info(f"Uploading file: {file_path} to container: {container_name}")
    connect_str = os.getenv('AZURE_STORAGE_CONNECTION_STRING')
    blob_service_client = BlobServiceClient.from_connection_string(connect_str)
    blob_client = blob_service_client.get_blob_client(container=container_name, blob=file_path)
    
    with open(file_path, "rb") as data:
        blob_client.upload_blob(data)
    return {"message": "File uploaded successfully"}

def translate_text(text, target_language):
    logging.info(f"Translating text: {text} to language: {target_language}")
    key = os.getenv("AZURE_TRANSLATOR_KEY")
    endpoint = os.getenv("AZURE_TRANSLATOR_ENDPOINT")
    credentials = CognitiveServicesCredentials(key)
    client = TextAnalyticsClient(endpoint=endpoint, credentials=credentials)
    
    response = client.translate(text, target_language)
    return response.translations[0].text

def get_job_insights():
    logging.info("Fetching job insights from LinkedIn API")
    headers = {
        'Authorization': 'Bearer YOUR_ACCESS_TOKEN',
    }
    response = requests.get('https://api.linkedin.com/v2/jobSearch', headers=headers)
    return response.json()